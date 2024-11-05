import logging
import os

import pandas as pd
import pypandoc
import requests
from docx import Document
from dotenv import load_dotenv
from jinja2 import Environment, FileSystemLoader

from utils import camel_case_to_capitalized, delete_empty_rows

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler("app.log"),
        logging.StreamHandler(),
    ],
)
logger = logging.getLogger(__name__)


def main():

    logger.info("Carregando constantes.")
    try:
        load_dotenv()

        API_KEY = os.getenv("API_KEY")
        BASE_URL = os.getenv("BASE_URL")
        PROTOCOL = 139

        URL = f"{BASE_URL}/{PROTOCOL}"
    except Exception as e:
        logger.error(e)
        exit(1)

    logger.info("Realizando requisição para API.")
    try:
        headers = {"Authorization": f"Bearer {API_KEY}"}
        response = requests.get(URL, headers=headers)

        response.raise_for_status()
        if response.status_code == 200:
            logger.info("Dados recebidos com sucesso.")
        else:
            raise (f"Failed to fetch data: {response.status_code} - {response.text}")

        vesgitio_list = response.json()["listaVestigio"]
        solicitacao = response.json()["solicitacaoOcorrencia"]

    except requests.exceptions.HTTPError as http_err:
        logger.error(f"HTTP error occurred: {http_err}")
        exit(1)
    except requests.exceptions.ConnectionError as conn_err:
        logger.error(f"Connection error occurred: {conn_err}")
        exit(1)
    except requests.exceptions.Timeout as timeout_err:
        logger.error(f"Timeout error occurred: {timeout_err}")
        exit(1)
    except requests.exceptions.RequestException as req_err:
        logger.error(f"An error occurred: {req_err}")
        exit(1)
    except Exception as e:
        logger.error(f"An unexpected error occurred: {e}")
        exit(1)

    try:
        logger.info("Gerando arquivo docx. [pydocx]")
        doc = Document()

        doc.add_heading(f"Protocolo número: {PROTOCOL}", level=0)

        doc.add_heading("Vestigios", level=1)

        for (
            i,
            v,
        ) in enumerate(vesgitio_list):
            doc.add_heading(f"Vestigio ID: {v["id"]}", level=2)
            doc.add_paragraph(v["resumo"], style="ListBullet")

            doc.add_heading(f"Tabela de dados {str(i+1)}", level=3)

            table_data = v["dadosVestigio"]
            table = doc.add_table(rows=len(table_data), cols=2)

            for i, (header, value) in enumerate(table_data.items(), start=0):

                if value != None:
                    cell = table.cell(i, 0)
                    cell.text = str(camel_case_to_capitalized(header))
                    cell = table.cell(i, 1)
                    cell.text = str(value)

            delete_empty_rows(table)
        logger.info("Salvando arquivo docx. [pydocx]")
        doc.save("doc_galileu.docx")

        # ODS
        logger.info("Carregando dataframe.")
        df = pd.json_normalize(solicitacao)

        logger.info("Convertendo datas.")
        df["dataInclusao"] = pd.to_datetime(df["dataInclusao"], unit="ms").dt.strftime("%d/%m/%Y %H:%M:%S")
        df["dataAcionamento"] = pd.to_datetime(df["dataAcionamento"], unit="ms").dt.strftime("%d/%m/%Y %H:%M:%S")
        df["dataHoraOcorrencia"] = pd.to_datetime(df["dataHoraOcorrencia"], unit="ms").dt.strftime("%d/%m/%Y %H:%M:%S")
        df["dataHoraAtendimentoPerito"] = pd.to_datetime(df["dataHoraAtendimentoPerito"], unit="ms").dt.strftime(
            "%d/%m/%Y %H:%M:%S"
        )
        df["dataHoraLiberacaoLocal"] = pd.to_datetime(df["dataHoraLiberacaoLocal"], unit="ms").dt.strftime(
            "%d/%m/%Y %H:%M:%S"
        )

        logger.info("Salvando arquivo ods.")
        df.to_excel("ods_galileu.ods", engine="odf")

    except Exception as e:
        logger.error(e)
        exit(1)

    try:
        logger.info("Iniciando processo com pandoc.")

        data = {"protocol": PROTOCOL, "vestigios": vesgitio_list}
        env = Environment(loader=FileSystemLoader("./templates"))
        template = env.get_template("template.html.j2")

        logger.info("Carregando template.")
        rendered_html = template.render(data)

        output_odt = "output_pandoc.odt"
        output_docx = "output_pandoc.docx"

        logger.info("Salvando arquivo odt.")
        pypandoc.convert_text(rendered_html, to="odt", format="html", outputfile=output_odt)

        logger.info("Salvando arquivo docx.")
        pypandoc.convert_text(rendered_html, to="docx", format="html", outputfile=output_docx)
    except Exception as e:
        logger.error(e)
        exit(1)

    logger.info("Processo finalizado com sucesso.")


if __name__ == "__main__":
    main()
